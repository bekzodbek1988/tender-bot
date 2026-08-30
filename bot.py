import os
import asyncio
import logging
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    ContextTypes,
    ConversationHandler,
    MessageHandler,
    filters,
)

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

SANA, BUYURTMACHI, LOT_NUMER = range(3)

COMPANY_NAME = "«Dobus Qurilish» МЧЖ"
DIRECTOR_NAME = "Рузиев Э. Б."

DOCUMENTS = {
    "1k_kafolat_xati": {
        "title": "Кафолат хати",
        "intro": "{COMPANY_NAME} шуни маълум қиладики, {SANA}даги {LOT_NUMER}-сонли тендерда энг яхши таклифларни танлашда иштирок этиб, товар етказиб бериш учун ёки иш ва хизматларни бажариш учун ишларни бажариш ва хизматларни кўрсатиш бўйича техник таклиф. Тендерда ғолиб деб топилган тақдирда, шартномада кўрсатиладиган ишларни бажариш ва хизматларни кўрсатиш муддатлари юзасидан қуйидаги мажбуриятларни ўз зиммасига олади ва тўлиқ кафолатлайди:",
        "items": [
            "Ишларни бошлаш муддати: Шартнома иккала томонлама имзоланиб, қонуний кучга киргандан ҳамда Давлат буюртмачиси томонидан олдиндан тўлов (аванс) маблағлари компаниямиз ҳисоб рақамига келиб тушган кундан бошлаб 3 (уч) банк кунидан кечиктирмасдан объектда ишларни амалда бошлашни;",
            "Шартнома шартларига амал қилиш: Ишлар ва хизматларни техник топшириқ ҳамда тасдиқланган Нуқсон далолатномасига мувофиқ, босқичма-босқич ва белгиланган муддатларда сифатли тарзда якунлашни;",
            "Техник база ва сафарбарлик: Ишларни муддатида тугатиш учун компаниямиз ихтиёридаги барча зарур муҳандислик-техник ходимларни, ишчи кучини ҳамда етарли миқдордаги асбоб-ускуналар ва технологик воситаларни объектга биринчи кунданоқ тўлиқ сафарбар этишни;"
        ],
        "outro": "Ушбу кафолат хати шартнома имзоланган кундан бошлаб, нуқсон далолатномасидаги барча ишлар тасдиқланган графикга асосан тўлиқ якунланиб, Буюртмачи томонидан якуний қабул қилиш-топшириш далолатномаси имзолангунга қадар ўз кучини сақлаб қолади."
    },
    "2k_tanishuv": {
        "title": "Танишув",
        "intro": "{COMPANY_NAME} шуни маълум қиладики, {SANA}даги {LOT_NUMER}-сонли тендерда энг яхши таклифларни танлашда иштирок этиб, товар етказиб бериш учун ёки иш ва хизматларни бажариш учун мазкур лот бўйича ўтказилаётган танлов ҳужжатлари, лойиҳа-смета ҳужжатлари ҳамда тасдиқланган Техник вазифа (ТЗ — Техническое задание) шартлари ва талаблари билан тўлиқ танишиб чиққанлигини маълум қилади.\n\nБиз мазкур танлов ҳужжатларидаги ва Техник вазифадаги (ТЗ) барча шартларни тўлиқ тушунган ҳолда, танловда ғолиб деб топилган тақдиртимизда қуйидагиларни кафолатлаймиз:",
        "items": [
            "Техник талабларга мувофиқлик: Объектдаги барча ишларни (хизматларни) Техник вазифада (ТЗ) кўрсатилган параметрлар, норматив стандартлар (ҚМҚ, ШНК ёки ГОСТ) ва давлат буюртмачисининг техник талабларига қатъий риоя қилган ҳолда амалга оширишни;",
            "Сифат кафолати: Ишларни бажаришда фойдаланиладиган барча материаллар, конструкциялар ва асбоб-ускуналарнинг сифатли бўлишини, уларнинг тегишли сифат ва мувофиқлик сертификатларига эга бўлишини ҳамда бажарилган ишларга шартномада белгиланган муддат давомида тўлиқ кафолат беришни;",
            "Муддатларга риоя қилиш: Барча ишлар ва хизматларни танлов шартлари ҳамда иккала томонлама тасдиқланадиган Календарь режа-графигида кўрсатилган муддатлардан кечиктирмасдан, ўз вақтида ва тўлиқ ҳажмда якунлашни;",
            "Малакали мутахассислар сафарбарлиги: Лойиҳани сифатли якунлаш учун объектга етарли даражада малакали ва тажрибали муҳандис-техник ходимларни (ИТР) ҳамда ишчи кучини жалб этишни."
        ],
        "outro": "Мазкур танловнинг Техник вазифасида (ТЗ) кўрсатилган талаблар юзасидан биз томондан ҳеч қандай эътирозлар мавжуд эмас ва биз ушбу шартлар асосида мажбуриятларимизни сидқидилдан бажаришга тайёрмиз."
    },
    "4k_usqunalar": {
        "title": "Ускуналар ҳақида маълумотнома",
        "intro": "{COMPANY_NAME} шуни маълум қиладики, {SANA}даги {LOT_NUMER}-сонли тендерда энг яхши таклифларни танлашда иштирок этиб, товар етказиб бериш учун ёки иш ва хизматларни бажариш учун барча зарурий жиҳозлар ва асбоб-ускуналарлар билан тулиқ таъминланган. Зарур хом ашё захираси мавжуд. Қуйида мавжуд асбоб-ускуналар руйхати келтирилиб утилган:",
        "items": [
            "Компрессор воздушный-Total TC1202411 24 L 1500W-1 дона",
            "Компрессор воздушный-Total TCS1120508, 50 litr-1 дона",
            "Болгарка- Crown 180 MM-2 дона",
            "Болгарка- Crown 125 MM-2 дона",
            "Перфоратор Ingco RH10506, 1050 Вт-1 дона",
            "Перфоратор Croon 32 1500W 220V/50Hz-1 дона",
            "Аккумуляторная дрель-шуруповерт LFINE 48B-2 дона",
            "Шуруповёрт Graff 48V-1 дона",
            "Сварочный аппарат Ураган MIG-350 (MIG / MMA)-1 дона",
            "Аппарат для сварки пластиковых труб АСПТ-2000А Ресанта арт. 65/136-2 дона",
            "Сварочный аппарат для пластиковых труб Ураган-2 дона",
            "Газовый сварочный аппарат-1 дона."
        ],
        "outro": ""
    },
    "5k_xodimlar": {
        "title": "Мутахассислар ва ходимлар ҳақида маълумотнома",
        "intro": "{COMPANY_NAME} шуни маълум қиладики, {SANA}даги {LOT_NUMER}-сонли тендер доирасида юклатилган иш ва хизматларни сифатли ҳамда ўз вақтида бажариш учун етарли даражадаги малакали муҳандис-техник ходимлар ва мутахассислар штатига эга:",
        "items": [
            "Бош муҳандис - 1 нафар (олий маълумотли, иш стажи 10 йилдан ортиқ);",
            "Прораб (Иш юритувчи) - 2 нафар (олий маълумотли, иш стажи 7 йил);",
            "Электргазпайвандчи - 4 нафар (малака сертификатига эга);",
            "Умумий қурилиш ишчилари - 12 нафар."
        ],
        "outro": "Барча ходимлар хавфсизлик техникаси ва меҳнат муҳофазаси бўйича тегишли йўриқномалардан ўтган."
    },
    "6k_ish_tajribasi": {
        "title": "Иш тажрибаси ҳақида маълумотнома",
        "intro": "{COMPANY_NAME} томонидан сўнгги йилларда муваффақиятли бажарилган ва фойдаланишга топширилган аналог объектлар ва қурилиш-монтаж ишлари рўйхати:",
        "items": [
            "2024-2025 йй. - Саноат ва фуқаролик объектларини мукаммал таъмирлаш ва реконструкция қилиш ишлари;",
            "2025-2026 йй. - Муҳандислик коммуникация тармоқлари ва биноларни пардозлаш ишлари."
        ],
        "outro": "Бажарилган барча ишлар буюртмачилар томонидан сифатли деб эътироф этилган ва эътирозлар мавжуд эмас."
    },
    "7k_moliyaviy": {
        "title": "Молиявий барқарорлик ҳақида маълумотнома",
        "intro": "{COMPANY_NAME} шуни маълум қиладики, {SANA}даги {LOT_NUMER}-сонли тендер шартларига мувофиқ корхонамиз молиявий жиҳатдан барқарор ва банк олдида тўлов муддати ўтган қарздорликларга эга эмас.",
        "items": [
            "Солиқ ва бошқа мажбурий тўловлар бўйича қарздорлик мавжуд эмас;",
            "Банк ҳисоб рақамларига картотека-2 ёки бошқа чекловлар қўйилмаган."
        ],
        "outro": ""
    },
    "8k_кафолат_сифат": {
        "title": "Сифат ва мувофиқлик кафолати",
        "intro": "{COMPANY_NAME} {SANA}даги {LOT_NUMER}-сонли лот бўйича бажариладиган барча ишлар ва фойдаланиладиган қурилиш материалларининг давлат стандартларига тўлиқ мослигини кафолатлайди.",
        "items": [
            "Фойдаланиладиган барча хом ашё ва материаллар сертификатланган;",
            "Бажарилган ишларга кафолат муддати 12 ойни ташкил этади."
        ],
        "outro": ""
    },
    "9k_мехнат_мухофазаси": {
        "title": "Меҳнат муҳофазаси ва техника хавфсизлиги кафолати",
        "intro": "{COMPANY_NAME} объектда ишларни амалга ошириш давомида меҳнат муҳофазаси ва техника хавфсизлиги қоидаларига қатъий риоя қилинишини таъминлашни ўз зиммасига олади.",
        "items": [
            "Ишчилар махсус кийим ва ҳимоя воситалари билан тўлиқ таъминланган;",
            "Экология ва атроф-муҳитни муҳофаза қилиш талабларига риоя қилинади."
        ],
        "outro": ""
    },
    "10k_субподряд": {
        "title": "Субподрядчиларни жалб этмаслик тўғрисида маълумотнома",
        "intro": "{COMPANY_NAME} {SANA}даги {LOT_NUMER}-сонли тендер бўйича барча ишларни учинчи шахсларни (субподрядчиларни) жалб этмаган ҳолда, ўз кучи ва техник имкониятлари билан бажаришни кафолатлайди.",
        "items": [],
        "outro": ""
    },
    "12k_тижорат_таклифи": {
        "title": "Тижорат таклифи кафолати",
        "intro": "{COMPANY_NAME} {SANA}даги {LOT_NUMER}-сонли лот бўйича тақдим этилаётган нарх таклифи асосли ва якуний эканлигини, шартнома тузиш жараёнида асоссиз оширилмаслигини кафолатлайди.",
        "items": [],
        "outro": ""
    }
}

def make_stamp_transparent(input_path, output_path):
    try:
        img = Image.open(input_path).convert("RGBA")
        datas = img.getdata()
        new_data = []
        for item in datas:
            if item[0] > 200 and item[1] > 200 and item[2] > 200:
                new_data.append((255, 255, 255, 0))
            else:
                new_data.append(item)
        img.putdata(new_data)
        img.save(output_path, "PNG")
    except Exception as e:
        print(f"Печат расмида хатолик: {e}")

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Ассалому алайкум! Ҳужжатларни шакллантириш учун **Санани** киритинг (масалан: `2026-йил «14» август`):", parse_mode="Markdown")
    return SANA

async def get_sana(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['sana'] = update.message.text
    await update.message.reply_text("Раҳмат! Энди **Буюртмачи номини** киритинг (масалан: `«ОКМК» АЖ`):", parse_mode="Markdown")
    return BUYURTMACHI

async def get_buyurtmachi(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['buyurtmachi'] = update.message.text
    await update.message.reply_text("Энди **Лот рақамини** киритинг (масалан: `8213557`):", parse_mode="Markdown")
    return LOT_NUMER

async def generate_and_send(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    lot_number = update.message.text
    sana = context.user_data.get('sana', '')
    buyurtmachi = context.user_data.get('buyurtmachi', '')

    await update.message.reply_text("⏳ Ҳужжатлар тайёрланмоқда...")

    base_dir = os.path.dirname(os.path.abspath(__file__))
    original_stamp = os.path.join(base_dir, "stamp.png")
    clean_stamp = os.path.join(base_dir, "stamp_transparent.png")

    if os.path.exists(original_stamp):
        make_stamp_transparent(original_stamp, clean_stamp)

    temp_dir = os.path.join(base_dir, f"temp_{lot_number}")
    os.makedirs(temp_dir, exist_ok=True)

    generated_files = []

    for key, doc_info in DOCUMENTS.items():
        doc = Document()
        
        # Саҳифа ҳошияларини сиқилтириш (1 саҳифага пўльно сиғиши учун)
        for section in doc.sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.0)

        # 4. ШРИФТ 14 pt ҚИЛИНДИ
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        # 1. Корхона шапкаси (Ташкилот бланкасида олиб ташланди)
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(1)
        p_logo.paragraph_format.keep_with_next = True
        r_logo = p_logo.add_run("ООО «DOBUS QURILISH»")
        r_logo.font.size = Pt(20)
        r_logo.bold = True

        p_req = doc.add_paragraph()
        p_req.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_req.paragraph_format.space_after = Pt(0)
        p_req.paragraph_format.keep_with_next = True
        r_req = p_req.add_run("г. Ташкент Мирабадский район массив Куйлюк-2, дом-9, кв-27. Тел: 71 290-93-78\np/с 2020 8000 2048 2684 7001, АТИБ «Ипотека банк» Миробад филиал, МФО: 00 420 ОКЭД: 43310, ИНН: 301 458 084")
        r_req.font.size = Pt(8.5)

        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_after = Pt(6)
        p_line.paragraph_format.keep_with_next = True
        r_line = p_line.add_run("_________________________________________________________________________________")
        r_line.bold = True
        r_line.font.size = Pt(8)

        # 2. Сана ва Буюртмачи ((танлов буюртмачиси) олиб ташланди)
        table_top = doc.add_table(rows=1, cols=2)
        table_top.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table_top)

        cell_date = table_top.cell(0, 0)
        cell_date.width = Cm(9.0)
        p_d = cell_date.paragraphs[0]
        p_d.paragraph_format.keep_with_next = True
        r_date = p_d.add_run(sana)
        r_date.font.size = Pt(14)
        r_date.bold = True
        r_date.underline = True

        cell_buy = table_top.cell(0, 1)
        cell_buy.width = Cm(9.0)
        p_buy = cell_buy.paragraphs[0]
        p_buy.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_buy.paragraph_format.keep_with_next = True
        r_buy = p_buy.add_run(f"«{buyurtmachi}»")
        r_buy.font.size = Pt(14)
        r_buy.bold = True

        # Сарлавҳа (14pt)
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(8)
        p_title.paragraph_format.space_after = Pt(8)
        p_title.paragraph_format.keep_with_next = True
        r_title = p_title.add_run(doc_info["title"])
        r_title.font.size = Pt(14)
        r_title.bold = True

        # Кириш матни (14pt)
        intro_text = doc_info["intro"].format(COMPANY_NAME=COMPANY_NAME, SANA=sana, LOT_NUMER=lot_number)
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_intro.paragraph_format.first_line_indent = Inches(0.4)
        p_intro.paragraph_format.space_after = Pt(4)
        p_intro.paragraph_format.keep_with_next = True
        r_intro = p_intro.add_run(intro_text)
        r_intro.font.size = Pt(14)

        # Бандлар (14pt)
        for idx, item in enumerate(doc_info["items"], 1):
            p_item = doc.add_paragraph()
            p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_item.paragraph_format.left_indent = Inches(0.2)
            p_item.paragraph_format.space_after = Pt(3)
            p_item.paragraph_format.keep_with_next = True
            r_item = p_item.add_run(f"{idx}. {item}")
            r_item.font.size = Pt(14)

        # Якуний матн (14pt)
        if doc_info["outro"]:
            p_outro = doc.add_paragraph()
            p_outro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_outro.paragraph_format.first_line_indent = Inches(0.4)
            p_outro.paragraph_format.space_before = Pt(4)
            p_outro.paragraph_format.space_after = Pt(6)
            p_outro.paragraph_format.keep_with_next = True
            r_outro = p_outro.add_run(doc_info["outro"])
            r_outro.font.size = Pt(14)

        # 3. Мўҳр ва Имзо (2-саҳифага ўтиб кетмаслиги таъминланди)
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(8)
        p_space.paragraph_format.keep_with_next = True

        table_sign = doc.add_table(rows=1, cols=3)
        table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table_sign)

        cell_left = table_sign.cell(0, 0)
        cell_left.width = Cm(6.0)
        cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_l = cell_left.paragraphs[0]
        p_l.paragraph_format.keep_with_next = True
        r_sig1 = p_l.add_run(f"{COMPANY_NAME} раҳбари")
        r_sig1.font.size = Pt(14)
        r_sig1.bold = True

        cell_middle = table_sign.cell(0, 1)
        cell_middle.width = Cm(6.0)
        cell_middle.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_stamp = cell_middle.paragraphs[0]
        p_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_stamp.paragraph_format.keep_with_next = True
        if os.path.exists(clean_stamp):
            p_stamp.add_run().add_picture(clean_stamp, width=Inches(1.8))

        cell_right = table_sign.cell(0, 2)
        cell_right.width = Cm(5.5)
        cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_sig_name = cell_right.paragraphs[0]
        p_sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sig_name.paragraph_format.keep_with_next = True
        r_sig2 = p_sig_name.add_run(f"________________ {DIRECTOR_NAME}")
        r_sig2.font.size = Pt(14)
        r_sig2.bold = True

        file_path = os.path.join(temp_dir, f"{key}_{lot_number}.docx")
        doc.save(file_path)
        generated_files.append(file_path)

    # 5. Барча ҳужжатларни (1k-12k) Телеграмга юбориш
    for file_p in generated_files:
        try:
            with open(file_p, 'rb') as doc_file:
                await update.message.reply_document(document=doc_file)
            await asyncio.sleep(0.5)
        except Exception as e:
            print(f"Файл юборишда хатолик: {e}")
        finally:
            if os.path.exists(file_p):
                os.remove(file_p)

    if os.path.exists(temp_dir):
        os.rmdir(temp_dir)

    await update.message.reply_text("✅ Барча ҳужжатлар (1k, 2k, 4k, 5k, 6k, 7k, 8k, 9k, 10k ва 12k) 14-шрифтда, 1 саҳифали форматда муваффақиятли тайёрланди ва юборилди!")
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Амал бекор қилинди.")
    return ConversationHandler.END

def main():
    TOKEN = "8996916069:AAFfxGbWY6YrK4f784ChJneTAg7tyuLoqW4"

    app = ApplicationBuilder().token(TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            SANA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_sana)],
            BUYURTMACHI: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_buyurtmachi)],
            LOT_NUMER: [MessageHandler(filters.TEXT & ~filters.COMMAND, generate_and_send)],
        },
        fallbacks=[CommandHandler('cancel', cancel)],
    )

    app.add_handler(conv_handler)
    print("Бот янгиланган параметрлар билан ишга тушди...")
    app.run_polling()

if __name__ == '__main__':
    main()
